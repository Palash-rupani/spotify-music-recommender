import pandas as pd
from docx import Document

# Load results
df_results = pd.read_csv("results_kmeans.csv")

# Find best
best_row = df_results.loc[df_results["silhouette"].idxmax()]
best_k = best_row["params"]

# Create document
doc = Document()
doc.add_heading("KMeans Clustering Results - Spotify Dataset", 0)

doc.add_heading("Results Table", level=1)

# Add table
table = doc.add_table(rows=1, cols=len(df_results.columns))
hdr_cells = table.rows[0].cells
for i, col in enumerate(df_results.columns):
    hdr_cells[i].text = col

for _, row in df_results.iterrows():
    row_cells = table.add_row().cells
    for i, col in enumerate(df_results.columns):
        row_cells[i].text = str(row[col])

doc.add_heading("Best Model", level=1)
doc.add_paragraph(
    "The best KMeans model was found with parameters: {}, "
    "achieving a silhouette score of {:.3f}, "
    "Davies–Bouldin index of {:.3f}, "
    "and Calinski–Harabasz score of {:.2f}."
    .format(best_k, best_row['silhouette'], best_row['davies_bouldin'], best_row['calinski_harabasz'])
)

doc.add_heading("Observations", level=1)
doc.add_paragraph(
    "Based on the silhouette score, the optimal number of clusters is 4. "
    "We observe that beyond k=5, the silhouette score declines, indicating weaker cluster separation. "
    "Thus, k=4 provides the best balance between compactness and separation."
)

# Save document
doc.save("KMeans_Results_Report.docx")
print("Report saved as KMeans_Results_Report.docx")
